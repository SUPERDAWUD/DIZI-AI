import pickle
import warnings
from sklearn.exceptions import InconsistentVersionWarning

# Silence scikit-learn model version mismatch warnings
warnings.filterwarnings("ignore", category=InconsistentVersionWarning)
# --- Load scikit-learn intent classifier ---
try:
    with open('intent_classifier.pkl', 'rb') as f:
        intent_clf_data = pickle.load(f)
    intent_vectorizer = intent_clf_data['vectorizer']
    intent_clf = intent_clf_data['clf']
except Exception:
    intent_clf = None
    intent_vectorizer = None
# --- Device selection for inference (CPU/GPU) ---
from flask import request as flask_request
inference_device = 'auto'

if 'app' in globals():
    @app.route('/api/set-device', methods=['POST'])
    @dev_mode_required
    def set_device():
        global inference_device
        data = flask_request.get_json()
        device = data.get('device', 'auto')
        inference_device = device
        # Update LocalGPT device
        if hasattr(local_gpt, 'set_device'):
            local_gpt.set_device(device)
        return jsonify({'device': device})
# --- Deep Learning Predictor ---
from predictor import predict
# --- Ask Gemini: wrapper for free_llm_generate (text) ---
def ask_gemini(prompt):
    """Call Gemini LLM for text generation."""
    return free_llm_generate(prompt, task='text')
import os
import json
import torch
import random
import difflib
import requests
import sympy
import re
import threading
from sympy import symbols, Eq, solve, simplify, expand, factor, diff, integrate
from dotenv import load_dotenv
try:
    from serpapi import GoogleSearch
except ImportError:  # serpapi package not installed
    GoogleSearch = None
from chatbot_model import NeuralNet
from utils import bag_of_words, tokenize, start_generator_api
import pytz
from datetime import datetime
from nltk.corpus import words as nltk_words
from bs4 import BeautifulSoup
import mimetypes
import io
import contextlib
import tempfile
import shutil
from flask import Flask, request, jsonify, render_template, redirect, url_for, session
from functools import wraps
from jsonschema import validate, ValidationError
import sys
import wikipedia
import google.generativeai as genai
# --- Local GPT and RAG ---
from gpt_local import LocalGPT
# --- Initialize Local GPT Model (load once) ---
local_gpt = LocalGPT('all')

# --- Flask App Initialization (ensure only one app instance) ---
if 'app' not in globals():
    app = Flask(__name__)
    app.secret_key = os.getenv('FLASK_SECRET_KEY', 'devsecret')

# --- Dev Mode Password (ensure defined) ---
DEV_MODE_PASSWORD = os.getenv('DEV_MODE_PASSWORD', 'supersecret')

# --- Dev Mode Decorator (ensure defined) ---
def dev_mode_required(f):
    @wraps(f)
    def decorated_function(*args, **kwargs):
        if not session.get('dev_mode'):
            return redirect(url_for('dev_login'))
        return f(*args, **kwargs)
    return decorated_function

# Ensure UPLOADS_DIR is defined at the top for all functions
UPLOADS_DIR = os.path.join(os.path.dirname(__file__), 'uploads')
VECTOR_INDEX_DIR = os.path.join(os.path.dirname(__file__), 'vector_index')
os.makedirs(VECTOR_INDEX_DIR, exist_ok=True)
_ST_MODEL = None

# Load keys
load_dotenv()
SERPAPI_KEY = os.getenv("SERPAPI_KEY")

# Aggregation mode: 'parallel' (gather/synthesize all) or 'priority' (pick best)
# Default to 'parallel' to look at everything concurrently and avoid prioritization.
AGGREGATION_MODE = os.getenv("AGGREGATION_MODE", "parallel").lower()
REASONING_MODE = os.getenv("REASONING_MODE", "auto").lower()  # off|auto|on
try:
    torch.set_float32_matmul_precision('high')
except Exception:
    pass

# Expose a simple API to toggle aggregation mode when Flask app context exists
if 'app' in globals():
    @app.route('/api/set-aggregation-mode', methods=['POST'])
    @dev_mode_required
    def set_aggregation_mode():
        global AGGREGATION_MODE
        data = flask_request.get_json(force=True, silent=True) or {}
        mode = (data.get('mode') or '').lower()
        if mode not in ('parallel', 'priority'):
            return jsonify({'ok': False, 'error': 'mode must be parallel or priority'}), 400
        AGGREGATION_MODE = mode
        return jsonify({'ok': True, 'mode': AGGREGATION_MODE})

# --- SerpAPI helper ---
def serpapi_search(params):
    """Search SerpAPI using the official client if available, otherwise via HTTP."""
    params = dict(params)
    if GoogleSearch is not None:
        return GoogleSearch(params).get_dict()
    if 'api_key' not in params:
        raise ValueError("SerpAPI 'api_key' is required")
    resp = requests.get("https://serpapi.com/search", params=params, timeout=5)
    resp.raise_for_status()
    return resp.json()

# Load intents
with open("data/intents.json", "r", encoding="utf-8") as f:
    intents = json.load(f)

# Load knowledge base
with open("data/knowledge_base.json", "r", encoding="utf-8") as f:
    kb_data = json.load(f)
    knowledge_base = {item["question"].lower(): item["answer"] for item in kb_data.get("questions", [])}

# Load model
FILE = "model.pth"
if torch.cuda.is_available():
    data = torch.load(FILE)
else:
    data = torch.load(FILE, map_location=torch.device("cpu"))
input_size = data["input_size"]
hidden_size = data["hidden_size"]
output_size = data["output_size"]
all_words = data["all_words"]
tags = data["tags"]
model_state = data["model_state"]

model = NeuralNet(input_size, hidden_size, output_size)
model.load_state_dict(model_state)
model.eval()

device = torch.device("cuda" if torch.cuda.is_available() else "cpu")
model.to(device)

USER_DATA_FILE = "user_data.json"
USER_QUESTIONS_FILE = "user_questions.json"
USER_LOCATION_FILE = "user_location.json"
user_data_lock = threading.Lock()
user_questions_lock = threading.Lock()
user_location_lock = threading.Lock()

try:
    nltk_words.words()
except LookupError:
    import nltk
    nltk.download('words')

# Cache English words for spell correction to avoid repeated lookups
ENGLISH_WORDS = set(nltk_words.words())

# Spell correction helper
def correct_word(word):
    """
    Correct a single word using a static cache of English words and difflib for fuzzy matching.
    Returns the corrected word if a close match is found, otherwise returns the original word.
    """
    # Use a static cache for faster lookup
    if word in ENGLISH_WORDS:
        return word
    matches = difflib.get_close_matches(word, ENGLISH_WORDS, n=1, cutoff=0.8)
    return matches[0] if matches else word

def correct_sentence(sentence):
    """
    Corrects each word in a sentence using correct_word, only if not already in the English word cache.
    Returns the corrected sentence as a string.
    """
    # Only correct words that are not in the cache (faster)
    return ' '.join(correct_word(w) if w not in ENGLISH_WORDS else w for w in sentence.split())

def atomic_write_json(data, filename):
    """
    Atomically writes JSON data to a file to avoid corruption.
    Writes to a temporary file and then moves it to the target filename.
    """
    # Write JSON atomically to avoid corruption
    dirpath = os.path.dirname(os.path.abspath(filename))
    with tempfile.NamedTemporaryFile('w', dir=dirpath, delete=False, encoding='utf-8') as tf:
        json.dump(data, tf, indent=2)
        tempname = tf.name
    shutil.move(tempname, filename)

def update_user_data(username):
    """
    Updates or creates user data in user_data.json.
    """
    try:
        if os.path.exists('user_data.json'):
            with open('user_data.json', 'r', encoding='utf-8') as f:
                data = json.load(f)
        else:
            data = {}
        if username not in data:
            data[username] = {}
        # Add more user fields as needed
        with open('user_data.json', 'w', encoding='utf-8') as f:
            json.dump(data, f, indent=2)
    except Exception as e:
        print(f"[ERROR] update_user_data: {e}")


def update_user_questions(username, question, max_questions=100):
    """
    Updates user_questions.json with the latest question for the user.
    """
    try:
        if os.path.exists('user_questions.json'):
            with open('user_questions.json', 'r', encoding='utf-8') as f:
                data = json.load(f)
        else:
            data = {}
        if username not in data:
            data[username] = []
        data[username].append({"question": question, "timestamp": time.time()})
        if len(data[username]) > max_questions:
            data[username] = data[username][-max_questions:]
        with open('user_questions.json', 'w', encoding='utf-8') as f:
            json.dump(data, f, indent=2)
    except Exception as e:
        print(f"[ERROR] update_user_questions: {e}")


def get_user_location(ip_address):
    """
    Gets user location from IP using multiple APIs for robustness.
    """
    try:
        import requests
        # Try ipinfo.io first
        try:
            resp = requests.get(f"https://ipinfo.io/{ip_address}/json", timeout=3)
            if resp.status_code == 200:
                data = resp.json()
                loc = data.get('loc', '').split(',')
                return {
                    'city': data.get('city'),
                    'region': data.get('region'),
                    'country': data.get('country'),
                    'lat': float(loc[0]) if len(loc) == 2 else None,
                    'lon': float(loc[1]) if len(loc) == 2 else None
                }
        except Exception:
            pass
        # Fallback: ip-api.com
        try:
            resp = requests.get(f"http://ip-api.com/json/{ip_address}", timeout=3)
            if resp.status_code == 200:
                data = resp.json()
                return {
                    'city': data.get('city'),
                    'region': data.get('regionName'),
                    'country': data.get('country'),
                    'lat': data.get('lat'),
                    'lon': data.get('lon')
                }
        except Exception:
            pass
    except Exception as e:
        print(f"[ERROR] get_user_location: {e}")
    return None


def update_user_location(username, ip_address):
    """
    Updates user_data.json with the user's latest location.
    """
    loc = get_user_location(ip_address)
    if not loc:
        return
    try:
        if os.path.exists('user_data.json'):
            with open('user_data.json', 'r', encoding='utf-8') as f:
                data = json.load(f)
        else:
            data = {}
        if username not in data:
            data[username] = {}
        data[username]['location'] = loc
        with open('user_data.json', 'w', encoding='utf-8') as f:
            json.dump(data, f, indent=2)
    except Exception as e:
        print(f"[ERROR] update_user_location: {e}")


def get_user_location_from_file(username):
    """
    Loads user location from user_data.json.
    """
    try:
        if os.path.exists('user_data.json'):
            with open('user_data.json', 'r', encoding='utf-8') as f:
                data = json.load(f)
            return data.get(username, {}).get('location')
    except Exception as e:
        print(f"[ERROR] get_user_location_from_file: {e}")
    return None


def get_local_time(location):
    try:
        import requests
        if not location or not location.get('lat') or not location.get('lon'):
            return "Sorry, I couldn't determine your local time."
        lat, lon = location['lat'], location['lon']
        resp = requests.get(f"https://timeapi.io/api/Time/current/coordinate?latitude={lat}&longitude={lon}", timeout=3)
        if resp.status_code == 200:
            data = resp.json()
            return f"Local time: {data.get('dateTime', 'Unknown')} ({data.get('timeZone', '')})"
    except Exception:
        pass
    return "Sorry, I couldn't determine your local time."


def get_weather(location):
    """
    Gets the weather for a given location dict (city, country) using SerpAPI/Google.
    Returns a string with the weather or an error message.
    """
    api_key = os.getenv("SERPAPI_KEY")
    city = location.get("city")
    country = location.get("country")
    if not api_key or not city:
        return "Weather API key or city missing."
    try:
        results = serpapi_search({
            "q": f"weather in {city} {country}",
            "api_key": api_key,
            "engine": "google",
        })
        answer_box = results.get("answer_box", {})
        if "temperature" in answer_box:
            return f"Weather in {city}: {answer_box['temperature']} {answer_box.get('unit', '')}, {answer_box.get('description', '')}"
        elif "snippet" in answer_box:
            return f"Weather: {answer_box['snippet']}"
        return "Weather info not found."
    except Exception as e:
        return f"Weather error: {e}"

# --- Enhanced: Robust intent/KB matching, grammar normalization, and speech upgrades ---
import difflib
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity

# Helper: Normalize grammar (basic)
def normalize_grammar(text):
    # Lowercase, remove extra spaces, basic punctuation fix
    text = text.lower().strip()
    text = re.sub(r'[^\w\s\.,?!]', '', text)
    text = re.sub(r'\s+', ' ', text)
    return text

# Helper: Semantic similarity (TF-IDF cosine)
def semantic_similarity(query, choices):
    try:
        from sklearn.feature_extraction.text import TfidfVectorizer
        vectorizer = TfidfVectorizer().fit([query] + choices)
        vectors = vectorizer.transform([query] + choices)
        sims = (vectors[0:1] @ vectors[1:].T).toarray()[0]
        idx = sims.argmax()
        return choices[idx], sims[idx]
    except Exception:
        return None, 0.0

# Improved intent response engine (fuzzy + semantic)
def get_intent_response(msg):
    # Use scikit-learn classifier if available
    if intent_clf and intent_vectorizer:
        X = intent_vectorizer.transform([msg])
        probs = intent_clf.predict_proba(X)[0]
        best_idx = probs.argmax()
        best_prob = probs[best_idx]
        best_tag = intent_clf.classes_[best_idx]
        if best_prob > 0.7:
            for intent in intents["intents"]:
                if best_tag == intent["tag"]:
                    response = random.choice(intent["responses"])
                    fun_prefixes = ["[AI] ", ""]
                    return random.choice(fun_prefixes) + response
        # fallback to old logic if not confident
    sentence = tokenize(msg)
    X = bag_of_words(sentence, all_words)
    X = X.reshape(1, X.shape[0])
    X = torch.from_numpy(X).to(device)
    output = model(X)
    _, predicted = torch.max(output, dim=1)
    tag = tags[predicted.item()]
    probs = torch.softmax(output, dim=1)
    prob = probs[0][predicted.item()]
    if prob.item() > 0.5:
        for intent in intents["intents"]:
            if tag == intent["tag"]:
                response = random.choice(intent["responses"])
                fun_prefixes = ["[AI] ", ""]
                return random.choice(fun_prefixes) + response
    # Fuzzy/semantic fallback
    all_patterns = []
    tag_map = {}
    for intent in intents["intents"]:
        for pattern in intent["patterns"]:
            all_patterns.append(pattern)
            tag_map[pattern] = intent["tag"]
    match, score = semantic_similarity(msg, all_patterns)
    if score > 0.5:
        tag = tag_map[match]
        for intent in intents["intents"]:
            if tag == intent["tag"]:
                return random.choice(intent["responses"])
    return None

# Improved KB fuzzy/semantic match
def get_from_knowledge_base(msg, threshold=0.7):
    msg = normalize_grammar(msg)
    kb_keys = list(knowledge_base.keys())
    match, score = semantic_similarity(msg, kb_keys)
    if score > threshold:
        return knowledge_base[match]
    # fallback to difflib
    match = difflib.get_close_matches(msg, kb_keys, n=1, cutoff=threshold)
    if match:
        return knowledge_base[match[0]]
    return None

# --- Improved speech extraction and generation ---
def extract_speech_request(msg):
    # Accepts poor grammar, e.g. "speech cat", "make speech cat funny", etc.
    msg = normalize_grammar(msg)
    if "speech" in msg:
        # Try to extract topic, tone, audience, length
        topic = None
        tone = None
        audience = None
        length = None
        # Extract topic (after 'speech' or 'about')
        match = re.search(r'speech(?: about| on| for| regarding)? ([\w\s,]+)', msg)
        if match:
            topic = match.group(1).strip()
        else:
            # fallback: first noun after 'speech'
            parts = msg.split('speech', 1)
            if len(parts) > 1:
                topic = parts[1].strip(' ,.')
        # Tone
        if 'funny' in msg:
            tone = 'funny'
        elif 'serious' in msg:
            tone = 'serious'
        elif 'motivational' in msg:
            tone = 'motivational'
        elif 'sad' in msg:
            tone = 'sad'
        # Audience
        aud_match = re.search(r'for ([\w\s]+)', msg)
        if aud_match:
            audience = aud_match.group(1).strip()
        # Length
        if 'short' in msg:
            length = 'short'
        elif 'long' in msg:
            length = 'long'
        return topic or msg, tone, audience, length
    return None, None, None, None

# Image generation (Replicate API)
def generate_image_from_prompt(prompt):
    import replicate
    replicate_token = os.getenv("REPLICATE_API_TOKEN")
    dalle_token = os.getenv("OPENAI_API_KEY")
    if not replicate_token and not dalle_token:
        return "No image generation API token is configured."
    try:
        # Parse style/model/quality hints from prompt
        style = None
        model = None
        quality = None
        # Example: "draw a cat in anime style, high quality, using SDXL"
        style_match = re.search(r"(anime|realistic|cartoon|sketch|oil painting|watercolor|pixel art|cyberpunk|fantasy|photorealistic)", prompt, re.I)
        if style_match:
            style = style_match.group(1).lower()
        model_match = re.search(r"(sdxl|dall[\.-]?e|dalle|stable diffusion|midjourney|anything v3|dreamshaper|openjourney)", prompt, re.I)
        if model_match:
            model = model_match.group(1).lower()
        quality_match = re.search(r"(high quality|ultra quality|4k|8k|hd|detailed|masterpiece)", prompt, re.I)
        if quality_match:
            quality = quality_match.group(1).lower()
        # Clean prompt for model
        clean_prompt = prompt
        for m in [style, model, quality]:
            if m:
                clean_prompt = re.sub(m, '', clean_prompt, flags=re.I)
        clean_prompt = clean_prompt.strip(', .')
        # Model selection logic
        if model and ("dall" in model or "openai" in model):
            # Use DALL·E via OpenAI API
            if not dalle_token:
                return "DALL·E API key is missing."
            import openai
            openai.api_key = dalle_token
            try:
                response = openai.Image.create(
                    prompt=clean_prompt + (f", {style}" if style else "") + (f", {quality}" if quality else ""),
                    n=1,
                    size="1024x1024"
                )
                url = response['data'][0]['url']
                return f"Here is your DALL·E image: {url}"
            except Exception as e:
                return f"DALL·E error: {e}"
        # Default: SDXL or other Replicate models
        replicate_model = "stability-ai/sdxl"
        if model:
            if "anime" in model or "anything" in model:
                replicate_model = "andite/anything-v4.0"
            elif "dreamshaper" in model:
                replicate_model = "lucataco/dreamshaper-v8"
            elif "openjourney" in model:
                replicate_model = "prompthero/openjourney"
            elif "realistic" in model:
                replicate_model = "stablediffusionapi/realistic-vision-v51"
            elif "midjourney" in model:
                replicate_model = "prompthero/openjourney"
            elif "dall" in model:
                # fallback to DALL·E above
                pass
        if style == "anime":
            replicate_model = "andite/anything-v4.0"
        elif style == "realistic":
            replicate_model = "stablediffusionapi/realistic-vision-v51"
        # Add quality to prompt
        full_prompt = clean_prompt
        if style:
            full_prompt += f", {style} style"
        if quality:
            full_prompt += f", {quality}"
        output = replicate.run(
            f"{replicate_model}",
            input={"prompt": full_prompt},
            api_token=replicate_token
        )
        if output and isinstance(output, list):
            return f"Here is your image: {output[0]}"
        elif isinstance(output, str):
            return f"Here is your image: {output}"
        return "Image generation failed."
    except Exception as e:
        return f"Image error: {e}"

# --- Optional: Image Captioning/Analysis (if file referenced) ---
def analyze_image_file(filename):
    import replicate
    replicate_token = os.getenv("REPLICATE_API_TOKEN")
    if not replicate_token:
        return "Replicate API token is missing."
    filepath = os.path.join(UPLOADS_DIR, filename)
    if not os.path.exists(filepath):
        return "File not found."
    try:
        # Use BLIP image captioning model
        output = replicate.run(
            "salesforce/blip:7c6b5cfae6e8e4e4b7e7b2e0b1e5e1b6b6e1b6e1b6e1b6e1b6e1b6e1b6e1b6",
            input={"image": open(filepath, "rb")},
            api_token=replicate_token
        )
        if output:
            return f"Image caption: {output}"
        return "Image analysis failed."
    except Exception as e:
        return f"Image analysis error: {e}"

# Image generation (Replicate API)
def generate_image_from_prompt(prompt):
    import replicate
    replicate_token = os.getenv("REPLICATE_API_TOKEN")
    if not replicate_token:
        return "Replicate API token is missing."
    try:
        # Use a popular, free, and public model: stability-ai/sdxl
        output = replicate.run(
            "stability-ai/sdxl:9fa24d7e8cf3e4c0b7e7b2e0b1e5e1b6b6e1b6e1b6e1b6e1b6e1b6e1b6e1b6",
            input={"prompt": prompt},
            api_token=replicate_token
        )
        if output and isinstance(output, list):
            return f"Here is your image: {output[0]}"
        elif isinstance(output, str):
            return f"Here is your image: {output}"
        return "Image generation failed."
    except Exception as e:
        return f"Image error: {e}"

def extract_language(msg):
    # Look for language keywords in the message
    languages = ['python', 'javascript', 'java', 'c++', 'c#', 'ruby', 'go', 'typescript', 'php', 'swift', 'kotlin', 'rust']
    for lang in languages:
        if lang in msg.lower():
            return lang
    return 'python'  # default

# Google fallback
def ask_google(query):
    # Only use Google fallback if the query is not handled by math, image, code, speech, intent, or knowledge base
    # This function should only be called as a true fallback
    if not SERPAPI_KEY:
        return "Google API key is missing."
    try:
        results = serpapi_search({
            "q": query,
            "api_key": SERPAPI_KEY,
            "engine": "google",
        })
        answer_box = results.get("answer_box", {})
        snippet = results.get("organic_results", [{}])[0].get("snippet", "")
        followup = None
        # Only generate a followup if the query is a question or informational
        is_question = any(q in query.lower() for q in ['what', 'who', 'when', 'where', 'why', 'how', '?'])
        if is_question:
            try:
                followup_resp = requests.post("http://localhost:5050/generate/followup", json={"topic": query})
                if followup_resp.ok:
                    followup = followup_resp.json().get("followup", None)
            except Exception:
                pass
        if "answer" in answer_box:
            response = f"Here’s what I found: {str(answer_box['answer'])}"
        elif snippet:
            response = f"Here’s a quick answer: {snippet}"
        else:
            response = "I searched high and low, but nothing turned up. Want to try rephrasing?"
        if followup:
            response += f"\n\n{followup}"
        return response
    except Exception as e:
        return f"Search error: {e}"

# Run code in a restricted namespace and capture output
def run_python_code(code):
    output = io.StringIO()
    try:
        with contextlib.redirect_stdout(output):
            exec(code, {'__builtins__': __builtins__}, {})
        return output.getvalue(), None
    except Exception as e:
        return None, str(e)

# Add support for code execution in multiple languages (Python, JavaScript, etc.)
def run_code(code, language='python'):
    import subprocess
    import tempfile
    if language == 'python':
        return run_python_code(code)
    elif language in ['js', 'javascript']:
        with tempfile.NamedTemporaryFile('w', suffix='.js', delete=False) as f:
            f.write(code)
            fname = f.name
        try:
            result = subprocess.run(['node', fname], capture_output=True, text=True, timeout=10)
            if result.returncode == 0:
                return result.stdout, None
            else:
                return None, result.stderr
        except Exception as e:
            return None, str(e)
        finally:
            os.remove(fname)
    elif language in ['bash', 'sh', 'shell']:
        with tempfile.NamedTemporaryFile('w', suffix='.sh', delete=False) as f:
            f.write(code)
            fname = f.name
        try:
            result = subprocess.run(['bash', fname], capture_output=True, text=True, timeout=10)
            if result.returncode == 0:
                return result.stdout, None
            else:
                return None, result.stderr
        except Exception as e:
            return None, str(e)
        finally:
            os.remove(fname)
    elif language in ['c', 'cpp', 'c++']:
        ext = '.c' if language == 'c' else '.cpp'
        with tempfile.NamedTemporaryFile('w', suffix=ext, delete=False) as f:
            f.write(code)
            fname = f.name
        exe = fname + '.exe'
        try:
            if language == 'c':
                compile_cmd = ['gcc', fname, '-o', exe]
            else:
                compile_cmd = ['g++', fname, '-o', exe]
            comp = subprocess.run(compile_cmd, capture_output=True, text=True)
            if comp.returncode != 0:
                return None, comp.stderr
            result = subprocess.run([exe], capture_output=True, text=True, timeout=10)
            if result.returncode == 0:
                return result.stdout, None
            else:
                return None, result.stderr
        except Exception as e:
            return None, str(e)
        finally:
            os.remove(fname)
            if os.path.exists(exe):
                os.remove(exe)
    elif language == 'java':
        with tempfile.NamedTemporaryFile('w', suffix='.java', delete=False) as f:
            f.write(code)
            fname = f.name
        classname = os.path.splitext(os.path.basename(fname))[0]
        try:
            comp = subprocess.run(['javac', fname], capture_output=True, text=True)
            if comp.returncode != 0:
                return None, comp.stderr
            result = subprocess.run(['java', '-cp', os.path.dirname(fname), classname], capture_output=True, text=True, timeout=10)
            if result.returncode == 0:
                return result.stdout, None
            else:
                return None, result.stderr
        except Exception as e:
            return None, str(e)
        finally:
            os.remove(fname)
            classfile = os.path.join(os.path.dirname(fname), classname + '.class')
            if os.path.exists(classfile):
                os.remove(classfile)
    return None, f"Language '{language}' not supported for execution."

# Math handler (basic & symbolic)
import signal

class TimeoutException(Exception):
    pass

def timeout_handler(signum, frame):
    raise TimeoutException()


def handle_math(msg):
    """
    Evaluate math expressions safely with SymPy and a short timeout.
    Supports basic arithmetic, derivative/integral/limit, solve, simplify/expand/factor,
    simple statistics (mean/std/variance), and small matrices.
    """
    from sympy import Matrix, Derivative, Integral, limit, summation, Product, Symbol
    x, y, z = symbols("x y z")
    original_msg = msg.strip()
    expr_in = original_msg.lower().replace("^", "**")
    for prefix in [
        "what is ", "whats ", "what's ", "calculate ", "solve ", "compute ",
        "evaluate ", "find ", "give me ", "tell me ", "can you solve ",
        "can you calculate ", "can you evaluate ", "can you compute "
    ]:
        if expr_in.startswith(prefix):
            expr_in = expr_in[len(prefix):].strip()
            break
    import concurrent.futures
    def math_eval(expr: str):
        try:
            if expr.startswith("matrix"):
                expr2 = expr.replace("matrix", "").strip()
                mats = [Matrix(eval(m)) for m in re.findall(r'\[\[.*?\]\]', expr2)]
                if len(mats) == 2 and ('*' in expr2 or 'x' in expr2):
                    return f"Matrix result: {mats[0] * mats[1]}"
                return f"Matrix: {mats[0]}" if mats else None
            if expr.startswith(("derivative", "differentiate")):
                e2 = expr.replace("derivative", "").replace("differentiate", "").strip()
                return f"Derivative: {diff(e2)}"
            if expr.startswith(("integral", "integrate")):
                e2 = expr.replace("integral", "").replace("integrate", "").strip()
                return f"Integral: {integrate(e2)}"
            if expr.startswith("limit"):
                parts = expr.replace("limit", "").strip().split(',')
                if len(parts) == 3:
                    e2, var, val = parts
                    return f"Limit: {limit(e2, Symbol(var.strip()), float(val.strip()))}"
            if expr.startswith(("sum", "summation")):
                e2 = expr.replace("summation", "sum").replace("sum", "").strip()
                m = re.match(r'(.+),\s*\((.+),\s*(.+),\s*(.+)\)', e2)
                if m:
                    f, v, a, b = m.groups()
                    return f"Summation: {summation(f, Symbol(v.strip()), int(a), int(b))}"
            if expr.startswith("product"):
                e2 = expr.replace("product", "").strip()
                m = re.match(r'(.+),\s*\((.+),\s*(.+),\s*(.+)\)', e2)
                if m:
                    f, v, a, b = m.groups()
                    return f"Product: {Product(f, Symbol(v.strip()), int(a), int(b)).doit()}"
            # Series expansion: series f(x) around a up to n
            if expr.startswith("series"):
                # Example: series sin(x) around 0 up to 5
                m = re.match(r'series\s+(.+)\s+around\s+([^\s]+)\s+up\s+to\s+(\d+)', expr)
                if m:
                    fexpr, around, n = m.groups()
                    var = Symbol('x')
                    try:
                        a = float(around)
                    except Exception:
                        a = 0
                    return str(sympy.series(sympy.sympify(fexpr), var, a, int(n)))
            # Solve a system: system x+y=2; x-y=0
            if expr.startswith("system") or (';' in expr and '=' in expr):
                eqs = []
                for part in re.split(r';|\n', expr.replace('system', '').strip()):
                    part = part.strip()
                    if not part:
                        continue
                    if '=' in part:
                        lhs, rhs = part.split('=',1)
                        eqs.append(Eq(sympy.sympify(lhs), sympy.sympify(rhs)))
                if eqs:
                    sols = solve(eqs)
                    return f"System solution: {sols}"
            if expr.startswith("mean"):
                nums = [float(n) for n in re.findall(r'-?\d+\.?\d*', expr)]
                return f"Mean: {sum(nums)/len(nums)}" if nums else None
            if expr.startswith(("std", "standard deviation")):
                nums = [float(n) for n in re.findall(r'-?\d+\.?\d*', expr)]
                if not nums:
                    return None
                mu = sum(nums)/len(nums)
                return f"Standard deviation: {(sum((u-mu)**2 for u in nums)/len(nums))**0.5}"
            if expr.startswith("variance"):
                nums = [float(n) for n in re.findall(r'-?\d+\.?\d*', expr)]
                if not nums:
                    return None
                mu = sum(nums)/len(nums)
                return f"Variance: {sum((u-mu)**2 for u in nums)/len(nums)}"
            if expr.startswith("solve"):
                e2 = expr.replace("solve", "").strip()
                if "=" in e2:
                    lhs, rhs = e2.split("=")
                    sol = solve(Eq(sympy.sympify(lhs), sympy.sympify(rhs)))
                else:
                    sol = solve(sympy.sympify(e2))
                return f"Solution: {sol}"
            if expr.startswith("simplify"):
                return f"Simplified: {simplify(expr.replace('simplify', '').strip())}"
            if expr.startswith("expand"):
                return f"Expanded: {expand(expr.replace('expand', '').strip())}"
            if expr.startswith("factor"):
                return f"Factored: {factor(expr.replace('factor', '').strip())}"
            if re.search(r"[\\+\\-\\*/]", expr) and not re.fullmatch(r"\\d+", expr):
                res = sympy.sympify(expr).evalf()
                if res == int(res):
                    res = int(res)
                return f"The answer is {res}"
            if re.search(r"[\\+\\-\\*/]", original_msg) and not re.fullmatch(r"\\d+", original_msg):
                res = sympy.sympify(original_msg).evalf()
                if res == int(res):
                    res = int(res)
                return f"The answer is {res}"
            return None
        except Exception as e:
            return f"Math error: {e}"
    with concurrent.futures.ThreadPoolExecutor(max_workers=1) as ex:
        fut = ex.submit(math_eval, expr_in)
        try:
            return fut.result(timeout=3)
        except concurrent.futures.TimeoutError:
            return "Math error: Calculation took too long. Please try a simpler question."
        except Exception as e:
            return f"Math error: {e}"
# --- File Redaction Feature ---
def redact_file_content(filename, redactions):
    filepath = os.path.join(UPLOADS_DIR, filename)
    if not os.path.exists(filepath):
        return None, "File not found."
    try:
        with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
            content = f.read()
        # Redact each word/phrase (case-insensitive)
        for phrase in redactions:
            pattern = re.compile(re.escape(phrase), re.IGNORECASE)
            content = pattern.sub('[REDACTED]', content)
        # Save redacted file
        redacted_name = f"redacted_{filename}"
        redacted_path = os.path.join(UPLOADS_DIR, redacted_name)
        with open(redacted_path, 'w', encoding='utf-8') as f:
            f.write(content)
        return redacted_name, None
    except Exception as e:
        return None, f"Redaction error: {e}"

# --- Conversation Memory (short-term and long-term) ---
CONVO_HISTORY_FILE = "conversation_history.json"
convo_history_lock = threading.Lock()

# Save each user message and bot reply with timestamp
import time

def save_conversation(username, user_msg, bot_reply):
    record = {
        "username": username,
        "timestamp": time.time(),
        "user_msg": user_msg,
        "bot_reply": bot_reply
    }
    with convo_history_lock:
        try:
            if os.path.exists(CONVO_HISTORY_FILE):
                with open(CONVO_HISTORY_FILE, "r", encoding="utf-8") as f:
                    history = json.load(f)
            else:
                history = []
        except Exception:
            history = []
        history.append(record)
        # Limit to last 1000 messages for short-term memory
        if len(history) > 1000:
            history = history[-1000:]
        try:
            atomic_write_json(history, CONVO_HISTORY_FILE)
        except Exception as e:
            print(f"[ERROR] Failed to write conversation history: {e}")

# Retrieve recent conversation for a user

def get_recent_conversation(username, n=20):
    with convo_history_lock:
        try:
            with open(CONVO_HISTORY_FILE, "r", encoding="utf-8") as f:
                history = json.load(f)
        except Exception:
            return []
        return [h for h in history if h["username"] == username][-n:]

# --- PDF, DOCX, Image OCR, and Document Q&A (Smarter) ---
import base64
from io import BytesIO

def extract_text_from_pdf(filepath):
    """
    Extracts text and tables from a PDF file using PyPDF2 and pandas (if possible).
    Returns the extracted text, tables, and error message (if any).
    """
    try:
        import PyPDF2
        with open(filepath, 'rb') as f:
            reader = PyPDF2.PdfReader(f)
            text = "\n".join(page.extract_text() or '' for page in reader.pages)
        # Try table extraction (very basic, for demonstration)
        tables = []
        try:
            import tabula
            tables = tabula.read_pdf(filepath, pages='all', multiple_tables=True)
            tables = [df.head(10).to_string() for df in tables if not df.empty]
        except Exception:
            pass
        return text, tables, None
    except Exception as e:
        return None, None, f"PDF extraction error: {e}"

def extract_text_from_docx(filepath):
    """
    Extracts text and tables from a DOCX file using python-docx.
    Returns the extracted text, tables, and error message (if any).
    """
    try:
        import docx
        doc = docx.Document(filepath)
        text = "\n".join([p.text for p in doc.paragraphs])
        # Table extraction
        tables = []
        for table in doc.tables:
            rows = []
            for row in table.rows:
                rows.append([cell.text for cell in row.cells])
            import pandas as pd
            df = pd.DataFrame(rows)
            tables.append(df.head(10).to_string(index=False, header=False))
        return text, tables, None
    except Exception as e:
        return None, None, f"DOCX extraction error: {e}"

def extract_text_from_image(filepath):
    """
    Extracts text from an image file using pytesseract OCR.
    Returns the extracted text and error message (if any).
    """
    try:
        from PIL import Image
        import pytesseract
        text = pytesseract.image_to_string(Image.open(filepath))
        return text, None
    except Exception as e:
        return None, f"Image OCR error: {e}"

def extract_text_from_xlsx(filepath):
    """
    Extracts text and tables from an Excel file using pandas.
    Returns the extracted text, tables, and error message (if any).
    """
    try:
        import pandas as pd
        xls = pd.ExcelFile(filepath)
        text = []
        tables = []
        for sheet in xls.sheet_names:
            df = xls.parse(sheet)
            tables.append(df.head(10).to_string())
            text.append(f"Sheet: {sheet}\n" + df.head(10).to_string())
        return "\n\n".join(text), tables, None
    except Exception as e:
        return None, None, f"Excel extraction error: {e}"

def detect_language_and_translate(text, target_lang='en'):
    """
    Detects language and translates to English if needed (using googletrans).
    Returns (translated_text, detected_lang, error)
    """
    try:
        from googletrans import Translator
        translator = Translator()
        detected = translator.detect(text)
        if detected.lang != target_lang:
            translated = translator.translate(text, dest=target_lang)
            return translated.text, detected.lang, None
        return text, detected.lang, None
    except Exception as e:
        return text, None, f"Translation error: {e}"

# --- Semantic Embedding Search (Sentence Transformers) ---
def get_semantic_embedding(text):
    """
    Returns a vector embedding for the given text using Sentence Transformers if available, else None.
    """
    try:
        from sentence_transformers import SentenceTransformer
        model = SentenceTransformer('all-MiniLM-L6-v2')
        return model.encode([text])[0]
    except Exception:
        return None

# --- Smart Semantic Q&A over Documents ---
def semantic_doc_qa(question, doc_text):
    """
    Answers a question over a document using semantic embedding search if possible, else TF-IDF.
    Returns the best matching answer span or None.
    """
    try:
        # Split doc into sentences/paragraphs
        import re
        chunks = re.split(r'(?<=[.!?])\s+', doc_text)
        # Try embedding-based search
        q_emb = get_semantic_embedding(question)
        if q_emb is not None:
            from sentence_transformers import util
            chunk_embs = [get_semantic_embedding(c) for c in chunks]
            sims = util.cos_sim([q_emb], chunk_embs)[0].tolist()
            idx = sims.index(max(sims))
            return chunks[idx]
        else:
            # Fallback: TF-IDF
            from sklearn.feature_extraction.text import TfidfVectorizer
            vectorizer = TfidfVectorizer().fit([question] + chunks)
            vectors = vectorizer.transform([question] + chunks)
            sims = (vectors[0:1] @ vectors[1:].T).toarray()[0]
            idx = sims.argmax()
            return chunks[idx]
    except Exception:
        return None

# --- LLM Fallback (OpenAI/Gemini/Local) ---
def llm_fallback(query, context=None):
    """
    Uses an LLM API (OpenAI, Gemini, or local) to answer the query, optionally with context.
    """
    api_key = os.getenv("OPENAI_API_KEY")
    if not api_key:
        return None
    try:
        import openai
        openai.api_key = api_key
        prompt = f"Context: {context}\n\nUser: {query}\nAI: " if context else query
        response = openai.ChatCompletion.create(
            model="gpt-3.5-turbo",
            messages=[{"role": "system", "content": "You are a helpful, expert AI assistant."},
                      {"role": "user", "content": prompt}]
        )
        return response['choices'][0]['message']['content'].strip()
    except Exception:
        return None

# --- Code Explanation and Inline Suggestion ---
def explain_code(code, language='python'):
    """
    Uses LLM to explain code and suggest improvements.
    """
    explanation = llm_fallback(f"Explain this {language} code and suggest improvements:\n{code}")
    return explanation or "Code explanation unavailable."

# --- Free LLM Integration: Gemini/HuggingFace for Code & Speech Generation ---
# --- Gemini Model Selection ---
GEMINI_MODELS = {
    'text': 'models/gemini-2.5-pro',  # Best for advanced text/chat
    'code': 'models/gemini-2.5-pro',  # Best for code generation
    'speech': 'models/gemini-2.5-pro',  # Speech/text generation
    'embedding': 'models/gemini-embedding-001',  # Embedding tasks
    'image': 'models/imagen-4.0-ultra-generate-001',  # Best for image generation
    'video': 'models/veo-3.0-generate-preview',  # Best for video generation
    'audio': 'models/gemini-2.5-flash-preview-native-audio-dialog',  # Audio/dialogue
}

def free_llm_generate(prompt, task='text', language='en'):
    """
    Uses free LLM APIs (Google Gemini, HuggingFace) for text/code/speech/image/video/audio generation.
    Returns the generated text or a clear error message.
    """
    gemini_key = os.getenv('GEMINI_API_KEY')
    if gemini_key:
        try:
            genai.configure(api_key=gemini_key)
            model_name = GEMINI_MODELS.get(task, GEMINI_MODELS['text'])
            model = genai.GenerativeModel(model_name)
            if task == 'image':
                response = model.generate_content(prompt)
            elif task == 'video':
                response = model.generate_content(prompt)
            elif task == 'audio':
                response = model.generate_content(prompt)
            elif task == 'embedding':
                response = model.embed_content(prompt)
            else:
                response = model.generate_content(prompt)
            # Handle different response types
            if hasattr(response, 'text'):
                return response.text.strip()
            elif hasattr(response, 'result'):
                return str(response.result)
            elif isinstance(response, dict):
                return str(response)
            else:
                return str(response)
        except Exception as e:
            return f"Gemini API error: {e}"
    hf_token = os.getenv('HF_API_TOKEN')
    if hf_token:
        try:
            start_time = time.time()
            headers = {"Authorization": f"Bearer {hf_token}"}
            payload = {"inputs": prompt}
            # Use code generation model for code, text model for speech
            if task == 'code':
                url = "https://api-inference.huggingface.co/models/bigcode/starcoder2-15b"
            else:
                url = "https://api-inference.huggingface.co/models/google/gemma-7b"
            resp = requests.post(url, headers=headers, json=payload, timeout=10)
            elapsed = time.time() - start_time
            if elapsed > 15:
                return f"HuggingFace API error: Response took too long ({elapsed:.1f}s). Please try again later."
            if resp.ok:
                result = resp.json()
                if isinstance(result, list) and 'generated_text' in result[0]:
                    return result[0]['generated_text'].strip()
                elif isinstance(result, dict) and 'generated_text' in result:
                    return result['generated_text'].strip()
                elif isinstance(result, list) and 'text' in result[0]:
                    return result[0]['text'].strip()
            return f"HuggingFace API error: {resp.text}"
        except Exception as e:
            return f"HuggingFace API error: {e}"
    return "No valid Gemini or HuggingFace API key found. Generation unavailable."

# Update all generation functions to use the correct task

def generate_code_text(prompt, language='en'):
    return free_llm_generate(prompt, task='code', language=language)

def generate_speech_text(topic, tone, audience, length):
    prompt = f"Write a {length} speech about {topic} for {audience} in a {tone} tone."
    return free_llm_generate(prompt, task='speech')

def generate_image_text(prompt):
    return free_llm_generate(prompt, task='image')

def generate_video_text(prompt):
    return free_llm_generate(prompt, task='video')

def generate_audio_text(prompt):
    return free_llm_generate(prompt, task='audio')

def generate_embedding_text(prompt):
    return free_llm_generate(prompt, task='embedding')

# --- Parallel Orchestration + Synthesis Helpers ---
def run_parallel(tasks, max_workers=None, timeout=12):
    """
    Run a list of (name, callable) in parallel threads and collect non-empty results.
    Returns a list of (name, result) preserving task names; exceptions return None.
    """
    import concurrent.futures
    results = []
    if max_workers is None:
        try:
            max_workers = int(os.getenv('MAX_WORKERS', '8'))
        except Exception:
            max_workers = 8
    with concurrent.futures.ThreadPoolExecutor(max_workers=max_workers) as ex:
        future_map = {ex.submit(func): name for name, func in tasks}
        for fut in concurrent.futures.as_completed(future_map, timeout=timeout):
            name = future_map.get(fut)
            try:
                val = fut.result(timeout=0)
                if val:
                    results.append((name, val))
            except Exception:
                # Skip failed tasks silently; other sources still contribute
                continue
    return results

def synthesize_from_sources(user_msg, sources, personality=None):
    """
    Ask an LLM to synthesize a single, accurate answer from multiple sources.
    If LLM synthesis is unavailable, fall back to a simple merged view.
    """
    # De-duplicate by name and trim very long entries
    seen = set()
    merged = []
    for name, val in sources:
        if not val:
            continue
        key = (name, str(val)[:100])
        if key in seen:
            continue
        seen.add(key)
        txt = val if isinstance(val, str) else json.dumps(val, ensure_ascii=False)
        merged.append((name, txt[:1200]))
    if not merged:
        return None
    bullet_lines = []
    for n, v in merged:
        bullet_lines.append(f"- {n}: {v}")
    persona_hint = ''
    if isinstance(personality, dict) and personality.get('preferred_tone'):
        persona_hint = f"Tone: {personality.get('preferred_tone')}\n"
    prompt = (
        "Synthesize the most accurate, concise answer using all evidence.\n"
        "Follow rules:\n"
        "- Prefer exact math/KB facts; avoid speculation.\n"
        "- If sources conflict, state the best-supported answer briefly.\n"
        "- Only include code if user asked for it; keep it minimal.\n"
        "- Be clear and helpful, in the user's language.\n\n"
        f"User message:\n{user_msg}\n\n"
        f"{persona_hint}Sources:\n" + "\n".join(bullet_lines) + "\n\nFinal answer:"
    )
    synthesized = free_llm_generate(prompt, task='text')
    if synthesized and not synthesized.lower().startswith(('no valid gemini', 'huggingface api error')):
        return synthesized.strip()
    # Fallback: simple merge
    return "\n\n".join([f"[{n}] {v}" for n, v in merged])

def ensure_relevance(query, answer, sources, threshold=0.22):
    """
    Ensure the answer is relevant to the query and the collected sources.
    If cosine-similarity is low, ask the LLM to refine for direct relevance.
    Returns possibly-refined answer.
    """
    try:
        from sklearn.feature_extraction.text import TfidfVectorizer
        corpus = [query, answer] + [str(v) for _, v in sources[:5]]
        vec = TfidfVectorizer().fit_transform(corpus)
        import numpy as np
        sim = (vec[0] @ vec[1].T).A[0][0]
        if sim < threshold:
            refine_prompt = (
                "Refine the following answer to directly and accurately address the user question.\n"
                "Only include information relevant to the question; drop unrelated text.\n\n"
                f"Question: {query}\n"
                f"Current answer: {answer}\n\n"
                "Refined answer:"
            )
            refined = free_llm_generate(refine_prompt, task='text')
            return refined or answer
    except Exception:
        pass
    return answer

def self_consistent_finalize(user_msg, sources, base_answer, samples=3):
    """
    Generate multiple candidates and pick the one most aligned with the question and sources.
    """
    candidates = [base_answer]
    try:
        styles = ["concise", "direct", "detailed"]
        for s in styles[:max(0, samples-1)]:
            prompt = (
                f"Rewrite the final answer in a {s} style while staying strictly on-topic.\n"
                f"Question: {user_msg}\n"
                f"Answer: {base_answer}\n"
                "Rewritten:"
            )
            cand = free_llm_generate(prompt, task='text')
            if cand:
                candidates.append(cand)
        # Score by TF-IDF similarity to question + sources
        from sklearn.feature_extraction.text import TfidfVectorizer
        texts = [user_msg] + [str(v) for _, v in sources[:8]]
        corpus = texts + candidates
        vec = TfidfVectorizer().fit_transform(corpus)
        import numpy as np
        q_vec = vec[0]
        src_vec = vec[1:1+len(texts)-1].mean(axis=0)
        best_idx = 0
        best_score = -1
        for i, _ in enumerate(candidates):
            c_vec = vec[len(texts)+i]
            score = (q_vec @ c_vec.T).A[0][0] + (src_vec @ c_vec.T).A[0][0]
            if score > best_score:
                best_score = score
                best_idx = i
        return candidates[best_idx]
    except Exception:
        return base_answer

# --- Personality Adaptation ---
def analyze_user_personality(username, msg):
    """
    Analyzes the user's message and recent history to infer personality traits (formality, humor, directness, tone).
    Updates user_data.json with detected traits.
    """
    traits = {
        'formality': None,
        'humor': None,
        'directness': None,
        'preferred_tone': None,
        'emoji_usage': None,
        'greeting_style': None
    }
    # Analyze current message
    msg_lc = msg.lower()
    if any(w in msg_lc for w in ['hi', 'hello', 'hey', 'greetings']):
        traits['greeting_style'] = 'casual'
    if any(e in msg for e in [':)', ':D', '😂', '😅', '😉', '😎', '🤣']):
        traits['humor'] = 'high'
        traits['emoji_usage'] = 'high'
    if any(w in msg_lc for w in ['please', 'could you', 'would you', 'kindly']):
        traits['formality'] = 'formal'
    if any(w in msg_lc for w in ['now', 'quick', 'fast', 'urgent']):
        traits['directness'] = 'direct'
    if any(w in msg_lc for w in ['funny', 'joke', 'laugh']):
        traits['preferred_tone'] = 'humorous'
    if any(w in msg_lc for w in ['serious', 'professional', 'business']):
        traits['preferred_tone'] = 'serious'
    # Analyze recent history
    recent = get_recent_conversation(username, n=10)
    humor_count = sum(1 for r in recent if any(e in r['user_msg'] for e in ['😂', '😅', '😉', '😎', '🤣']))
    formal_count = sum(1 for r in recent if any(w in r['user_msg'].lower() for w in ['please', 'kindly', 'would you']))
    if humor_count > 2:
        traits['humor'] = 'high'
    if formal_count > 2:
        traits['formality'] = 'formal'
    # Save traits to user_data.json
    try:
        with user_data_lock:
            if os.path.exists(USER_DATA_FILE):
                with open(USER_DATA_FILE, 'r', encoding='utf-8') as f:
                    data = json.load(f)
            else:
                data = {}
            if username not in data:
                data[username] = {}
            data[username]['personality'] = traits
            atomic_write_json(data, USER_DATA_FILE)
    except Exception as e:
        print(f"[ERROR] analyze_user_personality: {e}")
    return traits

def get_user_personality(username):
    """
    Loads personality traits for a user from user_data.json.
    """
    try:
        with user_data_lock:
            if os.path.exists(USER_DATA_FILE):
                with open(USER_DATA_FILE, 'r', encoding='utf-8') as f:
                    data = json.load(f)
                return data.get(username, {}).get('personality', {})
    except Exception as e:
        print(f"[ERROR] get_user_personality: {e}")
    return {}

# --- Deep Learning Personality Trait Inference ---
import torch.nn as nn
import torch.nn.functional as F

class PersonalityNet(nn.Module):
    def __init__(self, input_dim, hidden_dim, output_dim):
        super().__init__()
        self.fc1 = nn.Linear(input_dim, hidden_dim)
        self.fc2 = nn.Linear(hidden_dim, output_dim)
    def forward(self, x):
        x = F.relu(self.fc1(x))
        x = torch.sigmoid(self.fc2(x))
        return x

# Define trait names and model
PERSONALITY_TRAITS = ['formality', 'humor', 'directness', 'preferred_tone', 'emoji_usage', 'greeting_style']
PERSONALITY_TONE_MAP = ['neutral', 'humorous', 'serious']
PERSONALITY_INPUT_DIM = 6  # e.g., counts of keywords, emojis, etc.
PERSONALITY_HIDDEN_DIM = 16
PERSONALITY_OUTPUT_DIM = len(PERSONALITY_TRAITS)
personality_model = PersonalityNet(PERSONALITY_INPUT_DIM, PERSONALITY_HIDDEN_DIM, PERSONALITY_OUTPUT_DIM)
personality_model.eval()

# Helper: Extract features from user history
def extract_personality_features(history):
    features = [0]*PERSONALITY_INPUT_DIM
    for r in history:
        msg = r['user_msg'].lower()
        # Feature 0: formality
        if any(w in msg for w in ['please', 'kindly', 'would you']):
            features[0] += 1
        # Feature 1: humor
        if any(e in r['user_msg'] for e in ['😂', '😅', '😉', '😎', '🤣']):
            features[1] += 1
        # Feature 2: directness
        if any(w in msg for w in ['now', 'quick', 'fast', 'urgent']):
            features[2] += 1
        # Feature 3: preferred_tone
        if any(w in msg for w in ['funny', 'joke', 'laugh']):
            features[3] += 1
        if any(w in msg for w in ['serious', 'professional', 'business']):
            features[4] += 1
        # Feature 4: emoji usage
        if any(e in r['user_msg'] for e in [':)', ':D', '😂', '😅', '😉', '😎', '🤣']):
            features[5] += 1
    return torch.tensor(features, dtype=torch.float32)

# --- Enhanced Personality Analysis (Deep Learning) ---
def analyze_user_personality(username, msg):
    """
    Uses a neural network to infer personality traits from user history and current message.
    Updates user_data.json with detected traits.
    """
    recent = get_recent_conversation(username, n=10)
    features = extract_personality_features(recent + [{'user_msg': msg}])
    with torch.no_grad():
        output = personality_model(features)
    traits = {}
    for i, trait in enumerate(PERSONALITY_TRAITS):
        val = output[i].item()
        if trait == 'preferred_tone':
            if val < 0.33:
                traits[trait] = 'neutral'
            elif val < 0.66:
                traits[trait] = 'humorous'
            else:
                traits[trait] = 'serious'
        else:
            traits[trait] = 'high' if val > 0.5 else 'low'
    # Save traits to user_data.json
    try:
        with user_data_lock:
            if os.path.exists(USER_DATA_FILE):
                with open(USER_DATA_FILE, 'r', encoding='utf-8') as f:
                    data = json.load(f)
                if not isinstance(data, dict):
                    data = {}
            else:
                data = {}
            if username not in data or not isinstance(data[username], dict):
                data[username] = {}
            data[username]['personality'] = traits
            atomic_write_json(data, USER_DATA_FILE)
    except Exception as e:
        print(f"[ERROR] analyze_user_personality: {e}")
    return traits

# --- Enhanced get_response: synthesize from all sources and adapt personality ---
def get_response(msg, username=None, file_context=None, all_files=False, feedback=None, use_web=False, use_rag=False):
    # --- Intent check FIRST; in parallel mode treat as one of many sources ---
    msg_corrected = correct_sentence(msg)
    msg_norm = normalize_grammar(msg_corrected)
    sources = []
    intent_already_added = False
    intent_reply = get_intent_response(msg_corrected)
    if intent_reply:
        if AGGREGATION_MODE == 'parallel':
            sources.append(('intent', intent_reply))
            intent_already_added = True
        else:
            # Priority mode returns immediately for fast UX
            followup = None
            try:
                followup_resp = requests.post("http://localhost:5050/generate/followup", json={"topic": msg_corrected})
                if followup_resp.ok:
                    followup = followup_resp.json().get("followup", None)
            except Exception:
                pass
            return {"type": "intent", "content": intent_reply, "followup": followup}
    # --- Model/Personality selection per user ---
    import flask
    model = flask.session.get('model')
    personality = flask.session.get('personality')
    # If not set, try user_data.json
    if username and (not model or not personality):
        try:
            with open('user_data.json', 'r', encoding='utf-8') as f:
                data = json.load(f)
            user = data.get(username, {})
            if not model:
                model = user.get('preferred_model', 'all')
            if not personality:
                personality = user.get('preferred_personality', 'friendly')
        except Exception:
            model = model or 'all'
            personality = personality or 'friendly'
    model = model or 'all'
    personality = personality or 'friendly'
    # Set model/personality for local_gpt
    local_gpt.set_model(model, personality)
    # --- Custom Model Integration ---
    use_custom = (model == 'custom' or model == 'custom_llm')
    context = {
        'gemini_key': os.getenv('GEMINI_API_KEY'),
        'serpapi_key': os.getenv('SERPAPI_KEY'),
        'model': custom_model
    }
    if use_custom:
        # Use custom model for everything good
        result = custom_super_generate(msg_norm, context) or custom_everything_good(msg_norm, context)
        # Compose sources for unified handling
        sources.append(('custom_text', result.get('text')))
        sources.append(('custom_code', result.get('code')))
        sources.append(('custom_speech', result.get('speech')))
        sources.append(('custom_probability', result.get('probability')))
        if result.get('google'):
            sources.append(('custom_google', result.get('google')))
    # 0.1. Data Prediction (user can ask: predict [data])
    if msg_norm.lower().startswith('predict '):
        try:
            # Example: predict 1,2,3;4,5,6 (rows separated by ;)
            data_str = msg_norm[8:].strip()
            rows = [list(map(float, row.split(','))) for row in data_str.split(';')]
            pred = predict(rows)
            sources.append(('prediction', f"Prediction: {pred}"))
        except Exception as e:
            sources.append(('prediction', f"Prediction error: {e}"))
    # --- Personality adaptation ---
    personality = analyze_user_personality(username, msg) if username else {}
    # 0. Structured tool call\n    call = parse_tool_call(msg)\n    if call:\n        res = exec_tool(call['tool'], call.get('args', {}))\n        if res:\n            sources.append(('tool', json.dumps(res)))\n    # 0. Tool/plugin execution
    tool_result = call_tool_plugin(msg_norm)
    if tool_result:
        sources.append(('tool', tool_result))
    # 0.5. Location Q&A
    if any(q in msg_norm for q in ["where am i", "what's my location", "where are we", "my location", "where is my location"]):
        if username:
            loc = get_user_location_from_file(username)
            if loc:
                sources.append(('location', f"Your location: {loc.get('city', 'Unknown')}, {loc.get('region', '')}, {loc.get('country', '')} (lat: {loc.get('lat', '?')}, lon: {loc.get('lon', '?')})"))
            else:
                sources.append(('location', "Sorry, I couldn't determine your location."))
        else:
            sources.append(('location', "I need your username to look up your location."))
    # 0.6. Local time Q&A
    if any(q in msg_norm for q in ["what time is it", "what's the time", "local time", "current time", "my time", "what time is it here"]):
        if username:
            loc = get_user_location_from_file(username)
            if loc:
                sources.append(('time', get_local_time(loc)))
            else:
                sources.append(('time', "Sorry, I couldn't determine your location to get the time."))
        else:
            sources.append(('time', "I need your username to look up your local time."))
    # 1. Math
    math_reply = handle_math(msg_norm)
    if math_reply:
        sources.append(('math', math_reply))
    # 2. Image generation
    if any(w in msg_norm for w in ["draw", "generate image", "create image", "picture", "art"]):
        sources.append(('image', generate_image_from_prompt(msg)))
    # 3. Multi-file Q&A
    if all_files:
        answer = multi_file_semantic_qa(msg)
        if answer:
            sources.append(('files', answer))
    # 4. File Q&A (if file_context provided)
    if file_context:
        content, err = get_uploaded_file_content(file_context)
        if content:
            answer = semantic_doc_qa(msg, content)
            if answer:
                sources.append(('file', f"From your file: {answer}"))
            else:
                sources.append(('file', f"File content: {content[:500]}..."))
        elif err:
            sources.append(('file', err))
    # 5. Code generation (explicit request)
    if re.search(r'(generate|write|show|give) (me )?(code|javascript|python|java|c\+\+|c#|typescript|html|css)', msg_norm):
        lang = extract_language(msg_norm)
        code = generate_code_text(msg_norm, language=lang)
        if code:
            sources.append(('code', {'content': code, 'language': lang}))
        else:
            sources.append(('code', {'content': f"Sorry, code generation failed. Please specify the language and what you want.", 'language': lang}))
    # 6. Speech (explicit request, robust to vague/bad input)
    topic, tone, audience, length = extract_speech_request(msg_norm)
    if topic:
        # If topic is still None or too short, try using the whole message
        if not topic or len(topic.strip()) < 3:
            topic = msg_corrected.strip()
        if not topic or len(topic) < 3:
            sources.append(('speech', "Sorry, I couldn't understand your speech topic. Please rephrase."))
        else:
            speech = generate_speech_text(topic, tone or 'neutral', audience or 'general audience', length or 'short')
            if speech:
                sources.append(('speech', speech))
            else:
                sources.append(('speech', "Speech generation unavailable."))
    # 7. Intent/KB
    intent_reply = get_intent_response(msg_corrected)
    # In parallel mode, avoid early return here if we've already added intent
    if AGGREGATION_MODE == 'parallel' and intent_already_added:
        intent_reply = None
    # Only return intent immediately for greeting, goodbye, thanks, feeling, funny, creator, capabilities
    if intent_reply:
        greeting_tags = ["greeting", "goodbye", "thanks", "feeling", "funny", "creator", "capabilities"]
        tag = None
        for intent in intents["intents"]:
            if intent_reply.strip().replace("[AI says]","").replace("🤖","").replace("✨","").startswith(tuple(intent["responses"])):
                tag = intent["tag"]
                break
        if tag in greeting_tags:
            return intent_reply
    kb_reply = get_from_knowledge_base(msg_corrected)
    if kb_reply:
        sources.append(('kb', kb_reply))
    # Vector index RAG over uploads/data
    try:
        vec_hits = query_vector_index(msg_corrected, top_k=3)
        if vec_hits:
            ctx = '\n'.join([f"[{h['source']}] {h['chunk']}" for h in vec_hits])
            vec_prompt = f"Use the following retrieved context to answer the user.\n{ctx}\nQuestion: {msg_corrected}\nAnswer:"
            vec_ans = free_llm_generate(vec_prompt, task='text')
            if vec_ans:
                sources.append(('vector_kb', vec_ans))
    except Exception:
        pass

    # 7.5–11. Parallel external knowledge/LLM lookups
    def _task_gemini():
        try:
            return ask_gemini(msg_corrected)
        except Exception:
            return None

    def _task_gpt_rag():
        try:
            def retrieve_kb_context(query, kb=knowledge_base, top_n=3):
                matches = []
                for q, a in kb.items():
                    if q and query.lower() in q:
                        matches.append((q, a))
                if len(matches) < top_n:
                    for q, a in kb.items():
                        if len(matches) >= top_n:
                            break
                        if query.lower() in a.lower() and (q, a) not in matches:
                            matches.append((q, a))
                return matches[:top_n]
            kb_context = retrieve_kb_context(msg_corrected)
            if kb_context:
                context_str = '\n'.join([f"Q: {q}\nA: {a}" for q, a in kb_context])
                prompt = f"You are a helpful AI assistant. Use the following knowledge base to answer the user's question.\n{context_str}\nUser: {msg_corrected}\nAI:"
                return local_gpt.generate(prompt, max_length=180)
        except Exception:
            return None

    def _task_memory():
        try:
            if not username:
                return None
            recent = get_recent_conversation(username, n=5)
            if not recent:
                return None
            context = '\n'.join([f"User: {r['user_msg']}\nAI: {r['bot_reply']}" for r in recent])
            return free_llm_generate(f"{msg}\n\nContext:\n{context}", task='text')
        except Exception:
            return None

    def _task_llm():
        try:
            return free_llm_generate(msg_corrected, task='text')
        except Exception:
            return None

    def _task_google():
        try:
            return ask_google(msg_corrected)
        except Exception:
            return None

    def _task_wikipedia():
        try:
            if any(q in msg_norm for q in ['what', 'who', 'when', 'where', 'why', 'how', '?']):
                return f"Wikipedia: {wikipedia.summary(msg_corrected, sentences=2)}"
            return None
        except Exception:
            return None

    parallel_results = run_parallel([
        ('gemini', _task_gemini),
        ('gpt_rag', _task_gpt_rag),
        ('memory', _task_memory),
        ('llm', _task_llm),
        ('google', _task_google),
        ('wikipedia', _task_wikipedia),
    ], max_workers=6)
    for name, val in parallel_results:
        if val:
            sources.append((name, val))
    # --- Synthesize best answer ---
    followup = None
    best_lang = None
    if AGGREGATION_MODE == 'parallel':
        # Combine all sources and synthesize one final response
        enhanced_text = synthesize_from_sources(msg_corrected, sources, personality)
        if username and personality and enhanced_text:
            enhanced_text = adapt_response_style(enhanced_text, personality)
        best_type = 'multi'
    else:
        # Legacy priority selection
        priority = ['math', 'code', 'speech', 'intent', 'kb', 'file', 'files', 'memory', 'llm', 'tool', 'google', 'wikipedia', 'image', 'time', 'location']
        best_type = None
        best_content = None
        for p in priority:
            for src, val in sources:
                if src == p and val:
                    best_type = src
                    if src == 'code' and isinstance(val, dict):
                        best_content = val['content']
                        best_lang = val['language']
                    else:
                        best_content = val
                    break
            if best_content:
                break
        if not best_content and sources:
            best_type = 'multi'
            best_content = '\n\n'.join([f"[{src}] {val}" for src, val in sources if val])
        if username and personality and best_content:
            best_content = adapt_response_style(best_content, personality)
        # Optional: sentence-level enhancement (kept for legacy mode only)
        enhanced_sentences = []
        try:
            import re
            sentences = re.split(r'(?<=[.!?])\s+', str(best_content))
            for sent in sentences:
                if sent.strip():
                    enhanced = free_llm_generate(f"Enhance this sentence for clarity, creativity, and helpfulness: {sent}", task='text')
                    enhanced_sentences.append(enhanced)
            enhanced_text = '\n'.join(enhanced_sentences)
        except Exception:
            enhanced_text = best_content
    try:
        followup_resp = requests.post("http://localhost:5050/generate/followup", json={"topic": msg_corrected})
        if followup_resp.ok:
            raw_followup = followup_resp.json().get("followup", None)
            # Filter out irrelevant/generic follow-ups (e.g., Disqus comments)
            if raw_followup and isinstance(raw_followup, str):
                if ("Disqus" in raw_followup or "enable JavaScript" in raw_followup or "Comments" in raw_followup):
                    followup = None
                elif len(raw_followup.strip()) < 5:
                    followup = None
                else:
                    followup = raw_followup
    except Exception:
        pass
    if best_type == 'code':
        return {"type": "code", "content": enhanced_text, "language": best_lang, "followup": followup, "personality": personality}
    elif best_type:
        return {"type": best_type, "content": enhanced_text, "followup": followup, "personality": personality}
    else:
        return {"type": "none", "content": "Sorry, I didn't understand. Please try again with a clearer question or request.", "followup": followup}

# --- Response Style Adapter ---
def adapt_response_style(response, personality):
    """
    Modifies the response style/tone based on detected personality traits.
    """
    # Formality
    if personality.get('formality') == 'formal':
        response = f"Dear user, {response}".replace('!', '.')
    # Humor
    # (Removed emoji addition)
    # Directness
    if personality.get('directness') == 'direct':
        response = response.replace('I think', 'Here you go:').replace('Maybe', 'Definitely')
    # Emoji usage
    # (Removed emoji addition)
    # Greeting style
    if personality.get('greeting_style') == 'casual':
        response = "Hey! " + response
    # Serious tone
    if personality.get('preferred_tone') == 'serious':
        response = response.replace('!', '.')
    return response

# --- Structured Tools (Schemas + Executors) ---
TOOLS = {
    "calculator": {
        "schema": {
            "type": "object",
            "properties": {"expr": {"type": "string"}},
            "required": ["expr"]
        }
    },
    "wiki_summary": {
        "schema": {
            "type": "object",
            "properties": {"query": {"type": "string"}},
            "required": ["query"]
        }
    },
    "code_exec": {
        "schema": {
            "type": "object",
            "properties": {"language": {"type": "string"}, "code": {"type": "string"}},
            "required": ["language", "code"]
        }
    }
}


def exec_tool(name, args):
    try:
        # Schema validation
        schema = TOOLS.get(name, {}).get('schema')
        if schema is not None:
            validate(instance=args, schema=schema)
        if name == 'calculator':
            expr = args.get('expr','')
            val = handle_math(expr) or 'No result'
            return {"tool": name, "result": val}
        if name == 'wiki_summary':
            import wikipedia
            q = args.get('query','')
            return {"tool": name, "result": wikipedia.summary(q, sentences=2)}
        if name == 'code_exec':
            lang = args.get('language','python')
            code = args.get('code','')
            out, err = run_code(code, lang)
            return {"tool": name, "result": out or err or ''}
    except ValidationError as ve:
        return {"tool": name, "error": f"Invalid args: {ve.message}"}
    except Exception as e:
        return {"tool": name, "error": str(e)}
    return {"tool": name, "error": "Unknown tool"}
def parse_tool_call(text):
    """Parse inline JSON tool call: {"tool":"name","args":{...}}"""
    try:
        m = re.search(r"\{\s*\"tool\"\s*:\s*\"(.*?)\"[\s\S]*?\"args\"\s*:\s*(\{[\s\S]*\})\s*\}", text)
        if not m:
            return None
        tool = m.group(1)
        args = json.loads(m.group(2))
        if tool in TOOLS:
            return {"tool": tool, "args": args}
    except Exception:
        return None
    return None
# --- Tool/Plugin Execution (WolframAlpha, Wikipedia, Calculator, etc.) ---
def call_tool_plugin(query):
    """
    Calls external APIs/tools for advanced queries (math, facts, code search, etc.).
    Returns the tool result or None.
    """
    # Example: WolframAlpha
    if 'calculate' in query or 'math' in query or 'solve' in query:
        appid = os.getenv('WOLFRAMALPHA_APPID')
        if appid:
            try:
                import wolframalpha
                client = wolframalpha.Client(appid)
                res = client.query(query)
                answer = next(res.results).text
                return f"WolframAlpha: {answer}"
            except Exception:
                pass
    # Example: Wikipedia
    if 'wikipedia' in query or 'who is' in query or 'what is' in query:
        try:
            import wikipedia
            summary = wikipedia.summary(query, sentences=2)
            return f"Wikipedia: {summary}"
        except Exception:
            pass
    # Add more plugins/tools as needed
    return None

# --- Multi-Document/Context Q&A ---
def get_all_uploaded_files():
    """
    Returns a list of all uploaded file names in the uploads directory.
    """
    try:
        return [f for f in os.listdir(UPLOADS_DIR) if os.path.isfile(os.path.join(UPLOADS_DIR, f))]
    except Exception:
        return []

def multi_file_semantic_qa(question):
    """
    Searches all uploaded files for the best answer to the question using semantic_doc_qa.
    Returns the best answer and file name.
    """
    best_score = 0
    best_answer = None
    best_file = None
    for fname in get_all_uploaded_files():
        content, err = get_uploaded_file_content(fname)
        if content:
            answer = semantic_doc_qa(question, content)
            if answer:
                # Use TF-IDF similarity as a proxy for score
                _, score = semantic_similarity(question, [answer])
                if score > best_score:
                    best_score = score
                    best_answer = answer
                    best_file = fname
    if best_answer:
        return f"From {best_file}: {best_answer}"
    return None

# --- Uploaded File Content Extraction ---
def get_uploaded_file_content(filename):
    """
    Retrieves and returns the content of an uploaded file by filename.
    Supports PDF, DOCX, XLSX, and image files. Returns (content, error).
    """
    filepath = os.path.join(UPLOADS_DIR, filename)
    if not os.path.exists(filepath):
        return None, "File not found."
    try:
        if filename.lower().endswith('.pdf'):
            text, tables, err = extract_text_from_pdf(filepath)
            if err:
                return None, err
            return text, None
        elif filename.lower().endswith('.docx'):
            text, tables, err = extract_text_from_docx(filepath)
            if err:
                return None, err
            return text, None
        elif filename.lower().endswith('.xlsx'):
            text, tables, err = extract_text_from_xlsx(filepath)
            if err:
                return None, err
            return text, None
        elif filename.lower().endswith(('.png', '.jpg', '.jpeg', '.gif')):
            text, err = extract_text_from_image(filepath)
            if err:
                return None, err
            return text, None
        else:
            with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
                content = f.read()
            return content, None
    except Exception as e:
        return None, f"File extraction error: {e}"

# --- API Endpoints ---
@app.route('/api/chats')
def api_chats():
    try:
        if os.path.exists(CONVO_HISTORY_FILE):
            with open(CONVO_HISTORY_FILE, 'r', encoding='utf-8') as f:
                history = json.load(f)
        else:
            history = []
        # Group by session or user, or just show last N chats
        chats = []
        for i, entry in enumerate(history[-50:]):
            chats.append({
                'id': i,
                'title': entry.get('user_msg', '')[:30] or f'Chat {i+1}',
                'timestamp': entry.get('timestamp', '')
            })
        return jsonify(chats)
    except Exception as e:
        return jsonify([])

@app.route('/dev-login', methods=['GET', 'POST'])
def dev_login():
    error = None
    print('Session at /dev-login:', dict(session))  # Debug log
    if request.method == 'POST':
        password = request.form.get('password')
        if password == os.getenv('DEV_MODE_PASSWORD', 'supersecret'):
            session['dev_mode'] = True
            print('Login successful, session:', dict(session))  # Debug log
            return redirect(url_for('dev'))
        else:
            error = 'Incorrect password.'
            print('Login failed, session:', dict(session))  # Debug log
    return render_template('dev_login.html', error=error)

@app.route('/dev', methods=['GET', 'POST'])
def dev():
    print('Session at /dev:', dict(session))  # Debug log
    if not session.get('dev_mode'):
        print('Not logged in, redirecting to /dev-login')  # Debug log
        return redirect(url_for('dev_login'))
    print('Logged in, rendering dev_panel.html')  # Debug log
    return render_template('dev_panel.html')

@app.route('/dev/exec', methods=['POST'])
@dev_mode_required
def dev_exec():
    code = request.form.get('code')
    lang = request.form.get('language', 'python')
    # Allow advanced code execution, model retraining, etc.
    if lang == 'python':
        output, err = run_python_code(code)
        return jsonify({'output': output, 'error': err})
    # Add more advanced ML/hacking tools here as needed
    return jsonify({'error': 'Language not supported in dev mode.'})

@app.route('/dev/code-assistant')
@dev_mode_required
def code_assistant():
    return render_template('code_assistant.html')

@app.route('/api/code-assistant-chat', methods=['POST'])
def api_code_assistant_chat():
    data = request.get_json()
    message = data.get('message', '')
    code = data.get('code', '')
    language = data.get('language', 'python')
    # Use your AI backend for code help
    ai_reply = free_llm_generate(message, task='code')
    # Optionally execute code
    output = ''
    if code.strip():
        try:
            if language == 'python':
                import io, contextlib
                f = io.StringIO()
                with contextlib.redirect_stdout(f):
                    exec(code, {})
                output = f.getvalue()
            else:
                output = 'Code execution only supported for Python.'
        except Exception as e:
            output = f'Error: {e}'
    return jsonify({'reply': ai_reply, 'output': output})

# --- Self-improving/Auto-coding ability ---
def self_modify_code(new_code, target_file):
    """
    Allows the chatbot to write correct code to itself (self-improvement).
    Only available in dev mode for security.
    """
    try:
        with open(target_file, 'w', encoding='utf-8') as f:
            f.write(new_code)
        return True, None
    except Exception as e:
        return False, str(e)

@app.route('/dev/self-modify', methods=['POST'])
@dev_mode_required
def dev_self_modify():
    new_code = request.form.get('new_code')
    target_file = request.form.get('target_file')
    success, err = self_modify_code(new_code, target_file)
    return jsonify({'success': success, 'error': err})

@app.route('/dev/gemini_models')
@dev_mode_required
def dev_gemini_models():
    """
    Lists available Gemini models and their supported methods for the current API key.
    """
    gemini_key = os.getenv('GEMINI_API_KEY')
    if not gemini_key:
        return render_template('gemini_models.html', error='GEMINI_API_KEY not set in environment.', models=[])
    try:
        import google.generativeai as genai
        genai.configure(api_key=gemini_key)
        models = genai.list_models()
        model_list = []
        for m in models:
            model_list.append({
                'name': getattr(m, 'name', str(m)),
                'supported_generation_methods': getattr(m, 'supported_generation_methods', [])
            })
        return render_template('gemini_models.html', models=model_list)
    except Exception as e:
        return render_template('gemini_models.html', error=str(e), models=[])

@app.route('/api/gemini_image', methods=['POST'])
def api_gemini_image():
    data = request.get_json()
    prompt = data.get('prompt', '')
    result = generate_image_text(prompt)
    return jsonify({'result': result})

@app.route('/api/gemini_video', methods=['POST'])
def api_gemini_video():
    data = request.get_json()
    prompt = data.get('prompt', '')
    result = generate_video_text(prompt)
    return jsonify({'result': result})

@app.route('/api/gemini_audio', methods=['POST'])
def api_gemini_audio():
    data = request.get_json()
    prompt = data.get('prompt', '')
    result = generate_audio_text(prompt)
    return jsonify({'result': result})

@app.before_request
def auto_update_user_location():
    if request.endpoint not in ('static',):
        username = session.get('username')
        if username:
            ip = request.headers.get('X-Forwarded-For', request.remote_addr)
            update_user_location(username, ip)

# --- System Resource Monitoring (CPU, GPU, Energy) ---
def get_system_stats():
    """
    Returns a dict with current CPU, GPU, RAM, and energy usage stats.
    """
    stats = {}
    try:
        import psutil
        stats['cpu_percent'] = psutil.cpu_percent(interval=1)
        stats['ram_percent'] = psutil.virtual_memory().percent
        stats['ram_used'] = psutil.virtual_memory().used // (1024*1024)
        stats['ram_total'] = psutil.virtual_memory().total // (1024*1024)
    except Exception:
        stats['cpu_percent'] = stats['ram_percent'] = stats['ram_used'] = stats['ram_total'] = None
    # GPU stats (NVIDIA only)
    try:
        import GPUtil
        gpus = GPUtil.getGPUs()
        if gpus:
            stats['gpu_percent'] = gpus[0].load * 100
            stats['gpu_mem_used'] = gpus[0].memoryUsed
            stats['gpu_mem_total'] = gpus[0].memoryTotal
        else:
            stats['gpu_percent'] = stats['gpu_mem_used'] = stats['gpu_mem_total'] = None
    except Exception:
        stats['gpu_percent'] = stats['gpu_mem_used'] = stats['gpu_mem_total'] = None
    # Energy (battery)
    try:
        import psutil
        battery = psutil.sensors_battery()
        if battery:
            stats['battery_percent'] = battery.percent
            stats['power_plugged'] = battery.power_plugged
        else:
            stats['battery_percent'] = stats['power_plugged'] = None
    except Exception:
        stats['battery_percent'] = stats['power_plugged'] = None
    return stats

@app.route('/api/system_stats')
def api_system_stats():
    return jsonify(get_system_stats())

@app.route('/api/image_understand', methods=['POST'])
def api_image_understand():
    try:
        image_file = request.files.get('image')
        question = request.form.get('question', 'Describe this image')
        if not image_file:
            return jsonify({'error': 'No image uploaded.'})
        img_bytes = image_file.read()
        img_b64 = base64.b64encode(img_bytes).decode('utf-8')
        gemini_key = os.getenv('GEMINI_API_KEY')
        if not gemini_key:
            return jsonify({'error': 'Gemini API key missing.'})
        try:
            import google.generativeai as genai
            genai.configure(api_key=gemini_key)
            model = genai.GenerativeModel('gemini-pro-vision')
            # Gemini expects image as bytes, not base64, so pass img_bytes
            response = model.generate_content([question, img_bytes])
            result = response.text.strip() if hasattr(response, 'text') else str(response)
            return jsonify({'result': result})
        except Exception as e:
            return jsonify({'error': f'Gemini Vision error: {e}'})
    except Exception as e:
        return jsonify({'error': str(e)})

from custom_model.custom_llm import (
    load_custom_llm, generate_text as custom_generate_text,
    generate_code as custom_generate_code,
    generate_speech as custom_generate_speech,
    estimate_probability as custom_estimate_probability,
    everything_good as custom_everything_good,
    super_generate as custom_super_generate
)

# Load custom model (example, can be improved for device/model selection)
CUSTOM_MODEL_PATH = os.path.join('custom_model', 'custom_llm_weights.pth')
custom_model = load_custom_llm(CUSTOM_MODEL_PATH, device='auto')




# --- Web fetching and summarization ---
def extract_urls(text):
    try:
        return re.findall(r"https?://[^\s)]+", text)
    except Exception:
        return []

def fetch_url_content(url, timeout=8):
    try:
        headers = {"User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64)"}
        r = requests.get(url, headers=headers, timeout=timeout)
        r.raise_for_status()
        html = r.text
        soup = BeautifulSoup(html, 'html.parser')
        for tag in soup(["script", "style", "noscript"]):
            tag.decompose()
        text = "\n".join([t.strip() for t in soup.get_text("\n").splitlines() if t.strip()])
        if len(text) > 20000:
            text = text[:20000]
        title = soup.title.string.strip() if soup.title else url
        return {"url": url, "title": title, "text": text}
    except Exception as e:
        return {"url": url, "title": url, "text": f"[Fetch error] {e}"}

def summarize_text_block(text, question=None, max_tokens_hint=800):
    try:
        instr = (
            "Summarize the content clearly with key points and facts. "
            "If a question is provided, focus strictly on answering it based on this content."
        )
        if question:
            prompt = f"{instr}\n\nQuestion: {question}\n\nContent:\n{text}\n\nSummary:"
        else:
            prompt = f"{instr}\n\nContent:\n{text}\n\nSummary:"
        return free_llm_generate(prompt, task='text')
    except Exception:
        return None

def web_summarize_url(url, question=None):
    data = fetch_url_content(url)
    text = data.get('text') or ''
    if text.startswith('[Fetch error]'):
        return {"summary": text, "title": data.get('title', url), "url": url}
    summary = summarize_text_block(text, question)
    return {"summary": summary or text[:800], "title": data.get('title', url), "url": url}

def web_search_and_summarize(query, num_results=3):
    if not SERPAPI_KEY:
        return {"summary": "Google API key missing.", "results": []}
    try:
        results = serpapi_search({"q": query, "api_key": SERPAPI_KEY, "engine": "google"})
        links = []
        for item in results.get('organic_results', [])[:num_results]:
            link = item.get('link') or item.get('url')
            if link:
                links.append(link)
        summaries = []
        for link in links:
            summaries.append(web_summarize_url(link, question=query))
        bullets = []
        for s in summaries:
            bullets.append(f"- ({s['title']}) {s['summary']}")
        prompt = (
            "Using the multi-page summaries below, write a concise, accurate synthesis that directly answers the query. "
            "Include short inline citations [1], [2], ... mapping to the ordered sources.\n\n"
            f"Query: {query}\n\nSummaries by page:\n" + "\n".join(bullets) + "\n\nSynthesis:"
        )
        synthesis = free_llm_generate(prompt, task='text')
        cites = [s.get('url') for s in summaries]
        return {"summary": synthesis, "results": summaries, "citations": cites}
    except Exception as e:
        return {"summary": f"Search/summarization error: {e}", "results": []}

# --- Lightweight Vector Index over uploads/data ---
def _get_st_model():
    global _ST_MODEL
    if _ST_MODEL is None:
        try:
            from sentence_transformers import SentenceTransformer
            _ST_MODEL = SentenceTransformer('all-MiniLM-L6-v2')
        except Exception:
            _ST_MODEL = False
    return _ST_MODEL

def build_vector_index(sources=None, chunk_chars=600):
    """
    Build/update a simple vector index from text files in uploads/ and data/.
    Stores embeddings in npz and metadata in JSON.
    """
    try:
        import numpy as np
    except Exception:
        return False, 'numpy not available'
    st = _get_st_model()
    if not st:
        return False, 'SentenceTransformer unavailable'
    roots = sources or [UPLOADS_DIR, os.path.join(os.path.dirname(__file__), 'data')]
    records = []
    texts = []
    for root in roots:
        if not os.path.exists(root):
            continue
        for dirpath, _, filenames in os.walk(root):
            for fname in filenames:
                path = os.path.join(dirpath, fname)
                if not any(fname.lower().endswith(ext) for ext in ('.txt', '.md', '.json', '.py', '.html', '.csv', '.pdf', '.docx', '.xlsx')):
                    continue
                content, err = (get_uploaded_file_content(fname) if root == UPLOADS_DIR else (None, None))
                if root != UPLOADS_DIR:
                    try:
                        with open(path, 'r', encoding='utf-8', errors='ignore') as f:
                            content = f.read()
                    except Exception:
                        continue
                if not content:
                    continue
                for i in range(0, len(content), chunk_chars):
                    chunk = content[i:i+chunk_chars]
                    texts.append(chunk)
                    records.append({'source': os.path.relpath(path, os.path.dirname(__file__)), 'offset': i})
    if not texts:
        return False, 'No text to index'
    embs = st.encode(texts, show_progress_bar=False)
    import numpy as np
    try:
        import faiss
        index = faiss.IndexFlatIP(len(embs[0]))
        # Normalize for cosine via inner product
        X = embs / (np.linalg.norm(embs, axis=1, keepdims=True) + 1e-8)
        index.add(X.astype('float32'))
        faiss.write_index(index, os.path.join(VECTOR_INDEX_DIR, 'faiss.index'))
    except Exception:
        np.savez(os.path.join(VECTOR_INDEX_DIR, 'index.npz'), X=embs)
    with open(os.path.join(VECTOR_INDEX_DIR, 'meta.json'), 'w', encoding='utf-8') as f:
        json.dump({'records': records}, f, indent=2)
    return True, f'Indexed {len(texts)} chunks from {len(records)} records.'
def query_vector_index(query, top_k=5):
    try:
        import numpy as np
        st = _get_st_model()
        if not st:
            return []
        npz_path = os.path.join(VECTOR_INDEX_DIR, 'index.npz')
        meta_path = os.path.join(VECTOR_INDEX_DIR, 'meta.json')
        if not os.path.exists(npz_path) or not os.path.exists(meta_path):
            ok, _ = build_vector_index()
            if not ok:
                return []
        try:
            import faiss
            import numpy as np
            meta_path = os.path.join(VECTOR_INDEX_DIR, 'meta.json')
            index_path = os.path.join(VECTOR_INDEX_DIR, 'faiss.index')
            if os.path.exists(index_path):
                index = faiss.read_index(index_path)
                st = _get_st_model()
                q = st.encode([query])[0]
                qn = q / (np.linalg.norm(q) + 1e-8)
                D, I = index.search(np.array([qn], dtype='float32'), top_k)
                with open(meta_path, 'r', encoding='utf-8') as f:
                    meta = json.load(f)['records']
                results = []
                for i in I[0]:
                    if int(i) < 0 or int(i) >= len(meta):
                        continue
                    rec = meta[int(i)]
                    src = os.path.join(os.path.dirname(__file__), rec['source'])
                    chunk = ''
                    try:
                        with open(src, 'r', encoding='utf-8', errors='ignore') as f:
                            f.seek(rec['offset'])
                            chunk = f.read(600)
                    except Exception:
                        chunk = ''
                    results.append({'source': rec['source'], 'offset': rec['offset'], 'chunk': chunk})
                return results
        except Exception:
            pass
        import numpy as np
        data = np.load(npz_path)
        X = data['X']
        q = st.encode([query])[0]
        sims = (X @ q) / (np.linalg.norm(X, axis=1) * np.linalg.norm(q) + 1e-8)
        idxs = sims.argsort()[-top_k:][::-1]
        with open(meta_path, 'r', encoding='utf-8') as f:
            meta = json.load(f)['records']
        results = []
        for i in idxs:
            rec = meta[int(i)]
            src = os.path.join(os.path.dirname(__file__), rec['source'])
            chunk = ''
            try:
                with open(src, 'r', encoding='utf-8', errors='ignore') as f:
                    f.seek(rec['offset'])
                    chunk = f.read(600)
            except Exception:
                chunk = ''
            results.append({'source': rec['source'], 'offset': rec['offset'], 'chunk': chunk})
        return results
    except Exception:
        return []



















def generate_image_advanced(prompt, negative_prompt=None, seed=None, num_images=1):
    import replicate
    token = os.getenv("REPLICATE_API_TOKEN")
    if not token:
        return []
    try:
        args = {"prompt": prompt}
        if negative_prompt:
            args["negative_prompt"] = negative_prompt
        if seed is not None:
            try:
                args["seed"] = int(seed)
            except Exception:
                pass
        output = replicate.run(
            "stability-ai/sdxl:9fa24d7e8cf3e4c0b7e7b2e0b1e5e1b6b6e1b6e1b6e1b6e1b6e1b6e1b6e1b6",
            input=args,
            api_token=token
        )
        urls = []
        if isinstance(output, list):
            urls = [str(u) for u in output][:max(1, int(num_images or 1))]
        elif isinstance(output, str):
            urls = [output]
        return urls
    except Exception:
        return []
